from docx import Document

def read_docx(file_path):
    doc = Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs])

def write_docx(text, file_path):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(file_path)
