from docx import Document

def write_rules_to_docx(rules, output_path):
    doc = Document()
    doc.add_heading("Extracted Business Rules", level=1)
    for rule in rules:
        doc.add_paragraph(rule)
    doc.save(output_path)
